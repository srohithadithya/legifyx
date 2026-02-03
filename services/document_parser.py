"""
Legifyx Document Parser
Handles parsing of PDF, DOCX, TXT and image files
"""

import os
import re
from pathlib import Path
from typing import Optional, Tuple
import logging

logger = logging.getLogger(__name__)


class DocumentParser:
    """Parse various document formats into text"""
    
    SUPPORTED_FORMATS = ['.pdf', '.docx', '.doc', '.txt', '.jpg', '.jpeg', '.png']
    
    def __init__(self):
        self.ocr_service = None
    
    def parse(self, file_path: str) -> Tuple[str, dict]:
        """
        Parse document and extract text
        
        Args:
            file_path: Path to the document
            
        Returns:
            Tuple of (extracted_text, metadata)
        """
        path = Path(file_path)
        extension = path.suffix.lower()
        
        if extension not in self.SUPPORTED_FORMATS:
            raise ValueError(f"Unsupported format: {extension}")
        
        metadata = {
            "filename": path.name,
            "format": extension,
            "size_bytes": path.stat().st_size if path.exists() else 0
        }
        
        if extension == '.pdf':
            text = self._parse_pdf(file_path)
        elif extension in ['.docx', '.doc']:
            text = self._parse_docx(file_path)
        elif extension == '.txt':
            text = self._parse_txt(file_path)
        elif extension in ['.jpg', '.jpeg', '.png']:
            text = self._parse_image(file_path)
        else:
            text = ""
        
        metadata["char_count"] = len(text)
        metadata["word_count"] = len(text.split())
        
        return text, metadata
    
    def parse_bytes(self, file_bytes: bytes, filename: str) -> Tuple[str, dict]:
        """Parse document from bytes"""
        import tempfile
        
        extension = Path(filename).suffix.lower()
        
        with tempfile.NamedTemporaryFile(delete=False, suffix=extension) as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name
        
        try:
            return self.parse(tmp_path)
        finally:
            os.unlink(tmp_path)
    
    def _parse_pdf(self, file_path: str) -> str:
        """Parse PDF file"""
        try:
            import pdfplumber
            
            text_parts = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            
            return "\n\n".join(text_parts)
        
        except ImportError:
            try:
                from PyPDF2 import PdfReader
                
                reader = PdfReader(file_path)
                text_parts = []
                for page in reader.pages:
                    text_parts.append(page.extract_text())
                
                return "\n\n".join(text_parts)
            
            except ImportError:
                logger.error("No PDF library available. Install pdfplumber or PyPDF2")
                return ""
    
    def _parse_docx(self, file_path: str) -> str:
        """Parse DOCX file"""
        try:
            from docx import Document
            
            doc = Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text)
            
            return "\n\n".join(paragraphs)
        
        except ImportError:
            logger.error("python-docx not installed. Run: pip install python-docx")
            return ""
    
    def _parse_txt(self, file_path: str) -> str:
        """Parse TXT file"""
        encodings = ['utf-8', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        
        return ""
    
    def _parse_image(self, file_path: str) -> str:
        """Parse image using OCR"""
        try:
            import pytesseract
            from PIL import Image
            
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image, lang='eng+hin')
            
            return text
        
        except ImportError:
            logger.error("OCR libraries not available. Install pytesseract and Pillow")
            return ""
        except Exception as e:
            logger.error(f"OCR failed: {e}")
            return ""
    
    def validate_file(self, file_path: str) -> Tuple[bool, str]:
        """Validate file before parsing"""
        path = Path(file_path)
        
        if not path.exists():
            return False, "File does not exist"
        
        if path.suffix.lower() not in self.SUPPORTED_FORMATS:
            return False, f"Unsupported format. Supported: {', '.join(self.SUPPORTED_FORMATS)}"
        
        # Check file size (50MB limit)
        max_size = 50 * 1024 * 1024
        if path.stat().st_size > max_size:
            return False, "File too large. Maximum size is 50MB"
        
        return True, "Valid"
